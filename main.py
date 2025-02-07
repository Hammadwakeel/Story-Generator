import os
import uuid
import tempfile
import time
import re
import asyncio

from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

# Import the custom modules
from llm import get_llm
from prompt import story_request, generate_story, image_request, generate_image_prompt
from flux import generate_image

from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Create the FastAPI instance
app = FastAPI(
    title="Bedtime Story Generator API",
    description="API to generate a bedtime story with images and save as a docx document.",
    version="1.0.0"
)

# ---------------------------------------------------------------------------
# Pydantic model for validating the incoming story parameters
# ---------------------------------------------------------------------------
class StoryParams(BaseModel):
    Age: str
    Theme: str
    Pages: int
    Time: int
    Tone: str
    Setting: str
    Moral: str

# ---------------------------------------------------------------------------
# Helper functions (wrapped from your provided code)
# ---------------------------------------------------------------------------
def inference(llm_instance, story_params: dict) -> str:
    """
    Generates the story text from the LLM based on user parameters.
    """
    req = story_request(
        Age=story_params["Age"],
        Theme=story_params["Theme"],
        Pages=story_params["Pages"],
        Time=story_params["Time"],
        Tone=story_params["Tone"],
        Setting=story_params["Setting"],
        Moral=story_params["Moral"]
    )
    prompt_text = generate_story(req)
    print("\nGenerating story. Please wait...\n")
    response = llm_instance.invoke(prompt_text)
    return response.content

def parse_story_sections(story_text: str) -> list:
    """
    Parses the LLM-generated story into sections using markers enclosed in '**'.
    """
    pattern = r'\*\*(.*?)\*\*\s*'
    matches = list(re.finditer(pattern, story_text, flags=re.DOTALL))
    sections = []
    for i, match in enumerate(matches):
        marker = match.group(1).strip()
        start = match.end()
        end = matches[i+1].start() if (i+1) < len(matches) else len(story_text)
        content = story_text[start:end].strip()
        section_text = f"{marker}\n\n{content}" if content else marker
        sections.append(section_text)
    return sections

def generate_images_for_sections(sections: list, style: str = "sketch") -> list:
    """
    Generates an image for each story section.
    """
    image_paths = []
    for idx, section in enumerate(sections):
        print(f"Generating image for section {idx+1}...")
        img_req = image_request(style=style, bedtime_story_content=section)
        img_prompt = generate_image_prompt(img_req)
        image = generate_image(img_prompt)
        if image:
            temp_dir = tempfile.gettempdir()
            image_filename = os.path.join(temp_dir, f"section_{idx+1}_{uuid.uuid4().hex}.png")
            image.save(image_filename)
            image_paths.append(image_filename)
            print(f"Image for section {idx+1} saved as {image_filename}\n")
        else:
            print(f"Failed to generate image for section {idx+1}.\n")
            image_paths.append(None)
        time.sleep(1)  # Optional pause between image generations
    return image_paths

def save_story_to_docx(sections: list, image_paths: list, output_filename: str) -> None:
    """
    Saves the story sections and images into a formatted Word document.
    """
    document = Document()
    
    # If the first section is a title, use it as the document title.
    if sections and sections[0].startswith("Title:"):
        lines = sections[0].splitlines()
        title_line = lines[0].strip()  # e.g., "Title: The Amazing Adventure"
        title_text = title_line.replace("Title:", "").strip()
        document.core_properties.title = title_text
        document.add_heading(title_text, level=1)
        sections = sections[1:]
        if image_paths:
            image_paths = image_paths[1:]
    
    # Process remaining sections.
    for idx, section in enumerate(sections):
        lines = section.splitlines()
        if not lines:
            continue
        first_line = lines[0].strip()
        if any(first_line.startswith(marker) for marker in ["Opening Hook:", "Page", "Ending", "The End"]):
            document.add_heading(first_line, level=2)
            remaining_text = "\n".join(lines[1:]).strip()
            if remaining_text:
                document.add_paragraph(remaining_text)
        else:
            document.add_paragraph(section)
        
        # Insert the corresponding image (if available).
        if idx < len(image_paths) and image_paths[idx]:
            try:
                document.add_picture(image_paths[idx], width=Inches(4))
            except Exception as e:
                print(f"Error inserting image for section {idx+1}: {e}")
    
    document.save(output_filename)
    print(f"\n📖 Story saved to: {output_filename}")

def generate_story_docx(story_params: dict) -> str:
    """
    Complete pipeline:
      - Validates the API key
      - Generates the story text via the LLM
      - Parses the story into sections
      - Generates images for each section
      - Saves the complete story with images as a Word document
    Returns the filename of the saved document.
    """
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
    if not OPENAI_API_KEY:
        raise Exception("Error: OPENAI_API_KEY not found in environment variables.")
    
    llm_instance = get_llm(OPENAI_API_KEY)
    
    # Generate the story text from the LLM
    story_text = inference(llm_instance, story_params)
    print("\nStory generated successfully!\n")
    
    # Parse the story text into sections
    sections = parse_story_sections(story_text)
    
    # Generate images for each section
    image_paths = generate_images_for_sections(sections, style="sketch")
    
    # Create a unique filename for the docx file in a temporary directory
    output_filename = os.path.join(tempfile.gettempdir(), f"bedtime_story_{uuid.uuid4().hex}.docx")
    
    # Save the story and images to the Word document
    save_story_to_docx(sections, image_paths, output_filename=output_filename)
    
    return output_filename

# ---------------------------------------------------------------------------
# API Endpoints
# ---------------------------------------------------------------------------

@app.get("/", summary="Root Endpoint", description="Welcome message and API information.")
async def root():
    """
    Returns a welcome message and a link to the API documentation.
    """
    return {
        "message": "Welcome to the Bedtime Story Generator API!",
        "documentation": "/docs"
    }
    
@app.post(
    "/generate-story",
    summary="Generate a Bedtime Story Document",
    description="Generates a story with images based on input parameters and returns a docx file.",
    response_description="The generated Word document (.docx) file."
)
async def generate_story_endpoint(story_params: StoryParams):
    """
    API endpoint that runs the complete story-generation pipeline.
    It accepts story parameters as JSON, processes the story and images,
    and returns a downloadable Word document.
    """
    try:
        # Run the blocking story generation in a separate thread
        docx_file = await asyncio.to_thread(generate_story_docx, story_params.dict())
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    return FileResponse(
        path=docx_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(docx_file)
    )

@app.get("/health", summary="Health Check", description="Returns the API health status.")
async def health():
    return {"status": "ok"}

# ---------------------------------------------------------------------------
# Run the server with: uvicorn main:app --reload
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
